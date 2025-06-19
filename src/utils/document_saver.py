import os
from datetime import datetime
from docx import Document

class DocumentSaver:
    def save_txt(self, text, summary):
        """Save as text file"""
        filename = f"ocr_result_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
        try:
            with open(filename, 'w', encoding='utf-8-sig', errors='replace') as f:
                f.write(text)
                if summary:
                    f.write("\n\n=== Summary ===\n")
                    if 'summary' in summary:
                        f.write(summary['summary'])
        except UnicodeEncodeError:
            # UTF-8 failed, retry with CP949
            with open(filename, 'w', encoding='cp949', errors='replace') as f:
                f.write(text)
                if summary:
                    f.write("\n\n=== Summary ===\n")
                    if 'summary' in summary:
                        f.write(summary['summary'])
        return os.path.abspath(filename)

    def save_docx(self, text, summary):
        """Save as Word document"""
        filename = f"ocr_result_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        doc = Document()
        doc.add_paragraph(text)
        
        if summary:
            doc.add_heading('Summary', level=1)
            if 'summary' in summary:
                doc.add_paragraph(summary['summary'])
                
        doc.save(filename)
        return os.path.abspath(filename) 