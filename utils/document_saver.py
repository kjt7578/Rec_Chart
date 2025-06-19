import os
from datetime import datetime
from docx import Document
from typing import Dict, List

class DocumentSaver:
    def __init__(self, save_dir: str = "saved_documents"):
        """
        문서 저장 클래스
        save_dir: 저장할 디렉토리 경로
        """
        self.save_dir = save_dir
        os.makedirs(save_dir, exist_ok=True)

    def save_txt(self, text: str, summary: Dict[str, str], candidate_name: str = "Unknown") -> str:
        """
        Save text and summary as TXT file
        Returns: saved file path
        """
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{candidate_name}_{timestamp}.txt"
        filepath = os.path.join(self.save_dir, filename)

        try:
            with open(filepath, 'w', encoding='utf-8-sig', errors='replace') as f:
                f.write(f"Candidate: {candidate_name}\n")
                f.write(f"Interview Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
                
                f.write("=== Original Text ===\n")
                f.write(text)
                f.write("\n\n")
                
                f.write("=== Category Summary ===\n")
                for category, content in summary.items():
                    f.write(f"\n[{category}]\n")
                    f.write(content)
                    f.write("\n")
        except UnicodeEncodeError:
            # UTF-8 failed, retry with CP949
            with open(filepath, 'w', encoding='cp949', errors='replace') as f:
                f.write(f"Candidate: {candidate_name}\n")
                f.write(f"Interview Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
                
                f.write("=== Original Text ===\n")
                f.write(text)
                f.write("\n\n")
                
                f.write("=== Category Summary ===\n")
                for category, content in summary.items():
                    f.write(f"\n[{category}]\n")
                    f.write(content)
                    f.write("\n")

        return filepath

    def save_docx(self, text: str, summary: Dict[str, str], candidate_name: str = "Unknown") -> str:
        """
        Save text and summary as DOCX file
        Returns: saved file path
        """
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{candidate_name}_{timestamp}.docx"
        filepath = os.path.join(self.save_dir, filename)

        doc = Document()
        
        # Title
        doc.add_heading(f'Interview Record: {candidate_name}', 0)
        doc.add_paragraph(f'Interview Date: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        
        # Original text
        doc.add_heading('Original Text', level=1)
        doc.add_paragraph(text)
        
        # Category summary
        doc.add_heading('Category Summary', level=1)
        for category, content in summary.items():
            doc.add_heading(category, level=2)
            doc.add_paragraph(content)
        
        doc.save(filepath)
        return filepath 